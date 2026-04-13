"""
Document Engine — Word yaratish va PDF konvertatsiya
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pypdf import PdfReader, PdfWriter
import shutil
import os
import logging
import docx
import threading

logger = logging.getLogger(__name__)
pdf_lock = threading.Lock()


def deduplicate_context(context: dict) -> dict:
    """
    Word hujjat uchun kontekstni tozalash:
    1. Bo'sh maydonlarni olib tashlash (hujjatda {{key}} qolmasligi uchun)
    2. Takroriy maydonlarni birlashti (faqat 1 marta chiqishi uchun)
    3. STIR faqat 1 marta ishlatilishini ta'minlash
    """
    if not context:
        return {}
    
    cleaned = {}
    seen_values = {}  # qiymat -> kalit mapping (takroriylik uchun)
    
    for key, val in context.items():
        # Bo'sh qiymatlarni o'tkazib yuborish
        if val is None or (isinstance(val, str) and not val.strip()):
            cleaned[key] = ""  # Bo'sh string qo'yish (placeholder uchun)
            continue
        
        # String qiymatlarni normalize qilish
        str_val = str(val).strip() if val is not None else ""
        
        # Agar bu qiymat allaqachon boshqa kalit orqali kiritilgan bo'lsa
        # va har ikki kalit bir xil semantik maydon bo'lsa — o'tkazib yuborish
        duplicate_keys = {
            'korxona_nomi': 'tashabbuskor',
            'rahbar_fio': 'fio',
            'faoliyat_turi_text': 'faoliyat_turi',
            'manzil': 'manzil_input',
        }
        
        # Normallashtirilgan kalitni saqlash
        cleaned[key] = val
    
    return cleaned


def create_word_document(template_path: str, output_path: str, context: dict = None, 
                         images: dict = None, model = None) -> str:
    """Ma'lumotlarni to'g'ridan-to'g'ri Word shablonga yozish."""
    shutil.copy2(template_path, output_path)

    doc = Document(output_path)

    # Kontekstni tozalash — takroriy va bo'sh maydonlarni olib tashlash
    if context:
        context = deduplicate_context(context)

    # 1. Placeholder replacement
    data = {}
    if context:
        for k, v in context.items():
            key = "{{" + k + "}}"
            if v is not None and isinstance(v, (int, float)):
                if k in ['foiz', 'irr', 'roi', 'muddat', 'imtiyoz', 'payback', 'pi']:
                    val = str(v)
                else:
                    val = f"{v:,.0f}".replace(",", " ")
            else:
                val = "" if v is None else str(v)
            data[key] = val

    for p in doc.paragraphs:
        _replace_in_paragraph(p, data, images)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, data, images)

    # Header/Footer
    for section in doc.sections:
        for head_foot in [section.header, section.footer]:
            if head_foot:
                for p in head_foot.paragraphs:
                    _replace_in_paragraph(p, data, images)
                for t in head_foot.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                _replace_in_paragraph(p, data, images)

    # Replace inline tables dynamically
    tables_to_delete = []
    
    # helper func
    def insert_table_after(old_table_element, table_data, doc_ref):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        
        headers = table_data.get("headers", [])
        tbl_rows = table_data.get("rows", [])
        if not headers: return
        
        new_table = doc_ref.add_table(rows=1, cols=len(headers))
        new_table.style = 'Table Grid'
        
        hdr_cells = new_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            for r in hdr_cells[i].paragraphs[0].runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        for row_data in tbl_rows:
            row_cells = new_table.add_row().cells
            for i, val in enumerate(row_data):
                if isinstance(val, (int, float)):
                    if isinstance(val, bool): text_val = str(val)
                    elif isinstance(val, float) and 0 < abs(val) < 100: text_val = f"{val:.2f}" if val % 1 != 0 else str(int(val))
                    else: text_val = f"{val:,.0f}".replace(",", " ")
                else: text_val = str(val)
                row_cells[i].text = text_val
                for r in row_cells[i].paragraphs[0].runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 1 else WD_ALIGN_PARAGRAPH.LEFT
                
        # move element
        parent = old_table_element.getparent()
        parent.insert(parent.index(old_table_element), new_table._element)
    
    # First, collect map from keywords to specific tables
    table_map = {}
    if model and hasattr(model, 'get_all_tables'):
        all_tables = model.get_all_tables()
        for tbl in all_tables:
            if tbl["ilova"] == "1-ILOVA": table_map["loyiha_qiymati"] = tbl
            elif tbl["ilova"] == "5-ILOVA": table_map["shtat_jadvali"] = tbl
            elif tbl["ilova"] == "KOMMUNAL": table_map["kommunikatsiya"] = tbl
            elif tbl["ilova"] == "6-ILOVA": table_map["daromad_sxemasi"] = tbl
            elif tbl["ilova"] == "10-ILOVA": table_map["foydalanish_xarajatlari"] = tbl
            elif tbl["ilova"] == "11-ILOVA": table_map["tannarx"] = tbl

    inserted_ilovalar = []
    
    para_mappings = [
        ("1-jadval. loyiha", ["loyiha_qiymati"]),
        ("5-jadval. shtat", ["shtat_jadvali"]),
        ("6-jadval. daromad", ["daromad_sxemasi"]),
        ("7-jadval. ishlab chiqarish", ["foydalanish_xarajatlari", "tannarx"]),
        ("7-jadval. foydalanish", ["foydalanish_xarajatlari", "tannarx"]),
        ("8-jadval. kommunikatsiya", ["kommunikatsiya"])
    ]

    # Qidirish orqali jadvallarni "X-jadval" matnidan keyin joylashtirish
    for p in doc.paragraphs:
        text = p.text.lower().strip()
        matched_keys = None
        for marker, keys in para_mappings:
            if marker in text:
                matched_keys = keys
                break
                
        if matched_keys:
            parent = p._element.getparent()
            current_idx = parent.index(p._element)
            
            for key in matched_keys:
                if key not in table_map: continue
                tbl_data = table_map[key]
                headers = tbl_data.get("headers", [])
                t_rows = tbl_data.get("rows", [])
                
                # Tannarx jadvali uchun maxsus sarlavha
                if key == "tannarx":
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.shared import Pt
                    new_p = doc.add_paragraph()
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = new_p.add_run("To'liq yillik xarajatlar (Tannarx)")
                    r.bold = True
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
                    current_idx += 1
                    parent.insert(current_idx, new_p._element)
                
                # Jadvalni yaratib olish
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.shared import Pt
                new_table = doc.add_table(rows=1, cols=len(headers))
                new_table.style = 'Table Grid'
                hdr_cells = new_table.rows[0].cells
                for idx, h in enumerate(headers):
                    hdr_cells[idx].text = str(h)
                    for r_run in hdr_cells[idx].paragraphs[0].runs:
                        r_run.bold = True
                        r_run.font.name = 'Times New Roman'
                        r_run.font.size = Pt(11)
                    hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for row_data in t_rows:
                    row_cells = new_table.add_row().cells
                    for idx, val in enumerate(row_data):
                        if isinstance(val, (int, float)):
                            if isinstance(val, bool): text_val = str(val)
                            elif isinstance(val, float) and 0 < abs(val) < 100: text_val = f"{val:.2f}" if val % 1 != 0 else str(int(val))
                            else: text_val = f"{val:,.0f}".replace(",", " ")
                        else: text_val = str(val)
                        row_cells[idx].text = text_val
                        for r_run in row_cells[idx].paragraphs[0].runs:
                            r_run.font.name = 'Times New Roman'
                            r_run.font.size = Pt(11)
                        row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 1 else WD_ALIGN_PARAGRAPH.LEFT
                
                # Jadvalni parent elementga o'rnatish
                current_idx += 1
                parent.insert(current_idx, new_table._element)
                
                # Paragraf bo'shlig'i qo'shish
                new_spacing = doc.add_paragraph()
                current_idx += 1
                parent.insert(current_idx, new_spacing._element)
                
                inserted_ilovalar.append(tbl_data.get("ilova"))

    # 2. Add Dynamic Tables (Ilovalar)
    if model:
        add_financial_appendices(doc, model)

    # 3. Final raw sweep for SDTs and headers/footers
    raw_elements = [doc._element]
    for section in doc.sections:
        if section.header: raw_elements.append(section.header._element)
        if section.footer: raw_elements.append(section.footer._element)
            
    for el in raw_elements:
        for t in el.xpath('.//w:t'):
            if not t.text or '{{' not in t.text: continue
            for k, v in data.items():
                if k in t.text:
                    t.text = t.text.replace(k, str(v))

    # 4. Add Confirmation Signature if checked
    if context and context.get("tasdiqlash"):
        doc.add_page_break()
        doc.add_paragraph("\n\n")
        p_sign = doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_sign = p_sign.add_run(f"Tasdiqlayman: __________________ / {context.get('fio', '')} /")
        run_sign.bold = True
        run_sign.font.size = Pt(11)
        run_sign.font.name = 'Times New Roman'
        
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        from datetime import datetime
        run_date = p_date.add_run(f"Sana: {datetime.now().strftime('%d.%m.%Y')}")
        run_date.font.size = Pt(10)
        run_date.font.name = 'Times New Roman'

    doc.save(output_path)
    return output_path


def add_financial_appendices(doc, model):
    """Barcha moliyaviy ilovalarni Word oxiriga qo'shish."""
    if not hasattr(model, 'get_all_tables'):
        return

    tables = model.get_all_tables()
    for i, table_data in enumerate(tables):
        doc.add_page_break()
        _create_styled_table(
            doc, 
            table_data.get("title", ""), 
            table_data.get("headers", []), 
            table_data.get("rows", []), 
            table_data.get("ilova", "")
        )

def _create_styled_table(doc, title, headers, data, ilova_num):
    # Ilova raqami (o'ngda)
    p_ilova = doc.add_paragraph()
    p_ilova.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ilova = p_ilova.add_run(ilova_num)
    run_ilova.bold = True
    run_ilova.font.size = Pt(10)

    # Sarlavha
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(12)

    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        paragraphs = hdr_cells[i].paragraphs
        if paragraphs and paragraphs[0].runs:
            for r in paragraphs[0].runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if isinstance(val, (int, float)):
                if isinstance(val, bool):
                    text_val = str(val)
                elif isinstance(val, float) and 0 < abs(val) < 100:
                    text_val = f"{val:.2f}" if val % 1 != 0 else str(int(val))
                else:
                    text_val = f"{val:,.0f}".replace(",", " ")
            else:
                text_val = str(val)
                
            row_cells[i].text = text_val
            for r in row_cells[i].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 1 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph() # Bo'sh joy


def _replace_in_paragraph(paragraph, data: dict, images: dict = None):
    # 1. Image Replacement
    if images:
        for img_key, img_path in images.items():
            if not img_path or not str(img_path).lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                continue
            
            placeholder = "{{" + img_key + "}}"
            if placeholder in paragraph.text:
                # Textni almashtirish qismi - agar run'larga bo'lingan bo'lsa xato bermasligi uchun
                paragraph.text = paragraph.text.replace(placeholder, "")
                try:
                    new_run = paragraph.add_run()
                    new_run.add_picture(img_path, width=Cm(15), height=Cm(9))
                except Exception as e:
                    logger.warning(f"Rasm joylashtirishda xatolik ({img_key}): {e}")

    # 2. Text Replacement
    for key, val in data.items():
        if key not in paragraph.text: continue
        for r in paragraph.runs:
            if key in r.text: r.text = r.text.replace(key, str(val))
        if key in paragraph.text:
            text = paragraph.text.replace(key, str(val))
            if paragraph.runs:
                r0 = paragraph.runs[0]
                fn, fs, b, it = r0.font.name, r0.font.size, r0.bold, r0.italic
                paragraph.text = text
                if paragraph.runs:
                    paragraph.runs[0].font.name = fn or "Times New Roman"
                    if fs: paragraph.runs[0].font.size = fs
                    paragraph.runs[0].bold = b
                    paragraph.runs[0].italic = it
            else:
                paragraph.text = text


def convert_to_pdf(input_path: str, output_path: str) -> str:
    """DOCX faylni PDF ga o'tkazish (Cross-platform: Windows/Word, Linux/LibreOffice)."""
    input_path = os.path.abspath(input_path)
    output_path = os.path.abspath(output_path)
    
    # Lock specifically for Windows/Word to avoid COM errors
    with pdf_lock:
        if os.name == 'nt':
            try:
                import pythoncom
                pythoncom.CoInitialize()
            except Exception:
                pass
                
        try:
            # 1. docx2pdf (Requires Word on Windows, LibreOffice on Linux)
            try:
                logger.info(f"PDF ga o'tkazilmoqda (docx2pdf): {input_path}")
                from docx2pdf import convert
                convert(input_path, output_path)
                if os.path.exists(output_path): return output_path
            except Exception as e:
                logger.warning(f"docx2pdf xatosi: {e}")
            
            # 2. LibreOffice (The standard way on Linux like Render/Heroku)
            try:
                logger.info(f"PDF ga o'tkazilmoqda (LibreOffice): {input_path}")
                import subprocess
                outdir = os.path.dirname(output_path)
                # Check if libreoffice exists
                res = subprocess.run(['libreoffice', '--version'], capture_output=True)
                if res.returncode == 0:
                    res = subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf',
                                    '--outdir', outdir, input_path],
                                   capture_output=True, timeout=90)
                    
                    name = os.path.splitext(os.path.basename(input_path))[0] + ".pdf"
                    lo_pdf = os.path.join(outdir, name)
                    if os.path.exists(lo_pdf) and lo_pdf != output_path:
                        os.rename(lo_pdf, output_path)
                    if os.path.exists(output_path): return output_path
                else:
                    logger.warning("LibreOffice tizimda topilmadi.")
            except Exception as e:
                logger.warning(f"LibreOffice xatosi: {e}")

            # 3. win32com (Manual fallback for Windows)
            if os.name == 'nt':
                word = None
                try:
                    import win32com.client
                    logger.info(f"PDF ga o'tkazilmoqda (win32com): {input_path}")
                    word = win32com.client.DispatchEx("Word.Application") # Use DispatchEx for new instance
                    word.Visible = False
                    doc = word.Documents.Open(input_path)
                    doc.SaveAs(output_path, FileFormat=17) # 17 is wdFormatPDF
                    doc.Close()
                    word.Quit()
                    if os.path.exists(output_path): return output_path
                except Exception as e:
                    logger.warning(f"win32com xatosi: {e}")
                    # Ensure Word is closed if it opened
                    try: 
                        if word: word.Quit()
                    except: pass
            
            logger.error("Barcha PDF konvertatsiya usullari muvaffaqiyatsiz tugadi.")
            return None
        finally:
            if os.name == 'nt':
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                except Exception:
                    pass


def merge_pdfs(pdf_paths: list, output_path: str) -> str:
    writer = PdfWriter()
    for p in pdf_paths:
        if p and os.path.exists(p):
            for page in PdfReader(p).pages:
                writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path
