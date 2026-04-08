import docx

doc = docx.Document('template.docx')
for item in doc.paragraphs:
    if "jadval" in item.text.lower() or "shtat" in item.text.lower() or "kommunikatsiya" in item.text.lower() or "daromad" in item.text.lower():
        print("PARA:", item.text.strip())

for i, table in enumerate(doc.tables):
    found = False
    for row in table.rows:
        for cell in row.cells:
            if "loyiha" in cell.text.lower() or "shtat jadvali" in cell.text.lower() or "kommunikatsiya" in cell.text.lower() or "daromad" in cell.text.lower() or "ish haqi" in cell.text.lower() or "kommunal" in cell.text.lower() or "suv" in cell.text.lower():
                print(f"TABLE {i}: Found matching text ->", cell.text.strip()[:50])
                found = True
                break
        if found:
            break
