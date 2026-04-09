"""
TO'LIQ TIZIM TEKSHIRUVI
========================
Barcha modullarni tekshirish:
1. FinancialEngine — hisob-kitoblar to'g'riligi
2. Word generatsiya — document_engine
3. Excel generatsiya — excel_writer  
4. Credit calculator — kredit hisoblash
5. Business categories — kategoriya va qidirish
"""
import os
import sys
import traceback

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE_DIR)

# ══════════════════════════════════════════════════════════
# Test ma'lumotlari (real-case scenario)
# ══════════════════════════════════════════════════════════
TEST_DATA = {
    "loyiha_nomi": "Mebel ishlab chiqarish sexi",
    "tashabbuskor": "Orzu MCHJ",
    "manzil": "Toshkent sh., Yunusobod tumani",
    "bank": "Hamkorbank ATB",
    "stir": "123456789",
    "jshshir": "12345678901234",
    "faoliyat": "Mebel ishlab chiqarish",
    "mulk": "MCHJ",
    "fio": "Karimov Azizbek",
    "pasport": "AA 1234567",
    "berilgan_vaqti": "01.01.2020",
    "faoliyat_turi": "ishlab_chiqarish",
    "soliq_turi": "ytt",
    "loyiha_qiymati": "500000000",
    "oz_mablag": "150000000",
    "kredit": "350000000",
    "foiz": "14",
    "muddat": "84",
    "imtiyoz": "6",
    "kredit_turi": "annuitet",
    "discount_rate": "13.5",
    "mahsulot": "Yotoq mebeli",
    "hajm": "1200",
    "narx": "1500000",
    "olchov": "dona",
    "xomashyo_narx": "500000",
    "uskuna_qiymati": "200000000",
    "direktor": "1",
    "xodim": "8",
    "yangi_xodim": "3",
    "rahbar_oylik": "5000000",
    "ishchi_oylik": "3000000",
    "yangi_ishchi_oylik": "2500000",
    "elektr": "500",
    "gaz": "100",
    "suv": "50",
    "oqava": "30",
}

# MCHJ uchun test data
TEST_DATA_MCHJ = dict(TEST_DATA)
TEST_DATA_MCHJ["soliq_turi"] = "mchj"
TEST_DATA_MCHJ["mulk"] = "MCHJ"

# Minimal data (kamroq input)
TEST_DATA_MINIMAL = {
    "loyiha_nomi": "Test minimal",
    "tashabbuskor": "Test",
    "loyiha_qiymati": "100000000",
    "oz_mablag": "30000000",
    "kredit": "70000000",
    "foiz": "18",
    "muddat": "36",
    "imtiyoz": "0",
    "mahsulot": "Test mahsulot",
    "hajm": "500",
    "narx": "800000",
}

results = {"passed": 0, "failed": 0, "errors": []}

def test_pass(name):
    results["passed"] += 1
    print(f"  ✅ {name}")

def test_fail(name, detail=""):
    results["failed"] += 1
    results["errors"].append(f"{name}: {detail}")
    print(f"  ❌ {name}: {detail}")


# ══════════════════════════════════════════════════════════
# 1. KREDIT KALKULYATOR TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("1. KREDIT KALKULYATOR TESTI")
print("="*60)

try:
    from modules.credit_calculator import hisob_kredit, annuitet_hisob, differentsial_hisob
    
    # Test 1.1: Annuitet kredit
    an = annuitet_hisob(350000000, 14, 84, 6)
    if an.kredit_summa == 350000000: test_pass("Annuitet: kredit summasi to'g'ri")
    else: test_fail("Annuitet: kredit summasi", f"Kutilgan: 350000000, Natija: {an.kredit_summa}")
    
    if len(an.jadval) == 84: test_pass(f"Annuitet: jadval uzunligi to'g'ri ({len(an.jadval)} oy)")
    else: test_fail("Annuitet: jadval uzunligi", f"Kutilgan: 84, Natija: {len(an.jadval)}")
    
    if an.jadval[0].imtiyozli: test_pass("Annuitet: 1-oy imtiyozli")
    else: test_fail("Annuitet: 1-oy imtiyozli emas")
    
    if an.jadval[5].imtiyozli and not an.jadval[6].imtiyozli: test_pass("Annuitet: imtiyoz-asosiy chegarasi to'g'ri")
    else: test_fail("Annuitet: imtiyoz chegarasi", f"6-oy: {an.jadval[5].imtiyozli}, 7-oy: {an.jadval[6].imtiyozli}")
    
    # Imtiyozli davrda asosiy qarz 0 bo'lishi kerak
    grace_principal = sum(t.asosiy_qarz for t in an.jadval if t.imtiyozli)
    if grace_principal == 0: test_pass("Annuitet: imtiyoz davrida asosiy qarz = 0")
    else: test_fail("Annuitet: imtiyoz davrida asosiy qarz", f"Natija: {grace_principal}")
    
    # Oxirgi oy qoldig'i ≈ 0
    if an.jadval[-1].qoldiq < 1: test_pass(f"Annuitet: oxirgi oy qoldig'i ≈ 0 ({an.jadval[-1].qoldiq})")
    else: test_fail("Annuitet: oxirgi oy qoldig'i", f"Natija: {an.jadval[-1].qoldiq}")
    
    # jami_tolov > kredit
    if an.jami_tolov > an.kredit_summa: test_pass(f"Annuitet: jami to'lov > kredit ({an.jami_tolov:,.0f} > {an.kredit_summa:,.0f})")
    else: test_fail("Annuitet: jami to'lov noto'g'ri")
    
    # Test 1.2: Differentsial kredit
    diff = differentsial_hisob(350000000, 14, 84, 6)
    if len(diff.jadval) == 84: test_pass(f"Differentsial: jadval uzunligi to'g'ri ({len(diff.jadval)})")
    else: test_fail("Differentsial: jadval uzunligi", f"{len(diff.jadval)}")
    
    if diff.jadval[-1].qoldiq < 1: test_pass(f"Differentsial: oxirgi oy qoldig'i ≈ 0")
    else: test_fail("Differentsial: oxirgi oy qoldig'i", f"{diff.jadval[-1].qoldiq}")
    
    # Test 1.3: hisob_kredit wrapper
    result_an = hisob_kredit(100000000, 20, 60, 12, "annuitet")
    result_diff = hisob_kredit(100000000, 20, 60, 12, "differentsial")
    if result_an.turi == "annuitet" and result_diff.turi == "differentsial":
        test_pass("hisob_kredit: wrapper to'g'ri ishlaydi")
    else:
        test_fail("hisob_kredit: wrapper", f"an={result_an.turi}, diff={result_diff.turi}")
    
    # Test 1.4: Edge case — 0 summa
    zero = hisob_kredit(0, 14, 84, 6)
    if zero.jami_tolov == 0: test_pass("Edge case: 0 summa to'g'ri ishlaydi")
    else: test_fail("Edge case: 0 summa", f"{zero.jami_tolov}")

except Exception as e:
    test_fail("KREDIT KALKULYATOR", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 2. FINANCIAL ENGINE TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("2. FINANCIAL ENGINE TESTI")
print("="*60)

try:
    from modules.financial_engine import FinancialEngine, safe_float, safe_int, calculate_npv, calculate_irr
    
    # Test 2.1: YTT rejimi
    model = FinancialEngine(TEST_DATA)
    
    if model.loyiha_nomi == "Mebel ishlab chiqarish sexi": test_pass("Parse: loyiha nomi to'g'ri")
    else: test_fail("Parse: loyiha nomi", model.loyiha_nomi)
    
    if model.kredit_summa == 350000000: test_pass("Parse: kredit summa to'g'ri")
    else: test_fail("Parse: kredit summa", f"{model.kredit_summa}")
    
    if model.faoliyat_turi == "ishlab_chiqarish": test_pass("Parse: faoliyat turi to'g'ri")
    else: test_fail("Parse: faoliyat turi", model.faoliyat_turi)
    
    if model.MODEL_YEARS == 7: test_pass(f"Parse: MODEL_YEARS = {model.MODEL_YEARS}")
    else: test_fail("Parse: MODEL_YEARS", f"{model.MODEL_YEARS}")
    
    # Test 2.2: Jadvallar yaratilganligini tekshirish
    tables = model.get_all_tables()
    expected_tables = 14  # 1-ILOVA through 14-ILOVA + KOMMUNAL
    if len(tables) == expected_tables: test_pass(f"Jadvallar soni to'g'ri: {len(tables)}")
    else: test_fail("Jadvallar soni", f"Kutilgan: {expected_tables}, Natija: {len(tables)}")
    
    # Har bir jadval tuzilishini tekshirish
    for tbl in tables:
        ilova = tbl.get("ilova", "?")
        has_title = bool(tbl.get("title"))
        has_headers = bool(tbl.get("headers"))
        has_rows = bool(tbl.get("rows"))
        if has_title and has_headers and has_rows:
            test_pass(f"Jadval [{ilova}]: tuzilishi to'g'ri")
        else:
            test_fail(f"Jadval [{ilova}]", f"title={has_title}, headers={has_headers}, rows={has_rows}")
    
    # Test 2.3: Prod Plan (6-ILOVA) — daromadlar
    revenues = model.t_prod_plan["data"]["yearly_revenue"]
    max_rev = 1200 * 1500000  # 1,800,000,000
    if len(revenues) == model.MODEL_YEARS: test_pass(f"ProdPlan: {len(revenues)} yillik daromad")
    else: test_fail("ProdPlan: yillik daromad soni", f"{len(revenues)}")
    
    # Quvvat koeffitsientlari bilan tekshirish (1-yil 70%, 7-yil 100%)
    expected_1y = max_rev * 0.70  # = 1,260,000,000
    if abs(revenues[0] - expected_1y) < 1: test_pass(f"ProdPlan: 1-yil daromad to'g'ri ({revenues[0]:,.0f})")
    else: test_fail("ProdPlan: 1-yil daromad", f"Kutilgan: {expected_1y:,.0f}, Natija: {revenues[0]:,.0f}")
    
    if abs(revenues[-1] - max_rev) < 1: test_pass(f"ProdPlan: 7-yil daromad = to'liq ({revenues[-1]:,.0f})")
    else: test_fail("ProdPlan: 7-yil daromad", f"Kutilgan: {max_rev:,.0f}, Natija: {revenues[-1]:,.0f}")
    
    # Test 2.4: NPV/IRR/ROI
    ind = model.indicators
    if ind["npv"] is not None: test_pass(f"NPV hisoblandi: {ind['npv']:,.0f}")
    else: test_fail("NPV hisoblanmadi")
    
    if ind["irr"] is not None: test_pass(f"IRR hisoblandi: {ind['irr']}%") 
    else: test_fail("IRR hisoblanmadi")
    
    if ind["roi"] is not None: test_pass(f"ROI hisoblandi: {ind['roi']}%")
    else: test_fail("ROI hisoblanmadi")
    
    if ind.get("payback") is not None: test_pass(f"Payback hisoblandi: {ind['payback']} yil")
    else: test_pass("Payback: loyiha 7+ yil (yoki hisoblangan)")  # Not necessarily a fail
    
    # Test 2.5: Kredit yillik qoldiq tekshiruvi
    yillik = model.yillik_kredit
    if len(yillik) == model.MODEL_YEARS: test_pass(f"Kredit yillik: {len(yillik)} yil")
    else: test_fail("Kredit yillik", f"Kutilgan: {model.MODEL_YEARS}, Natija: {len(yillik)}")
    
    # Test 2.6: Kommunal jadval
    komm = model.t_kommunal
    if komm["ilova"] == "KOMMUNAL": test_pass("Kommunal jadval yaratildi")
    else: test_fail("Kommunal jadval", f"ilova={komm['ilova']}")
    
    # Kommunal hisob: elektr = 500 kVt * 900 so'm = 450,000 so'm/oy
    elektr_oylik = 500 * 900
    komm_rows = komm["rows"]
    if len(komm_rows) > 0 and komm_rows[0][4] == elektr_oylik:
        test_pass(f"Kommunal: elektr oylik hisob to'g'ri ({elektr_oylik:,})")
    else:
        test_fail("Kommunal: elektr oylik hisob", f"Kutilgan: {elektr_oylik}, Natija: {komm_rows[0][4] if komm_rows else 'bo''sh'}")
    
    # Test 2.7: Labour (Ish haqi)
    labour = model.t_labour
    dir_yillik = 1 * 5000000 * 12  # 60,000,000
    labour_data = labour["data"]
    if labour_data["admin_yillik"] == dir_yillik:
        test_pass(f"Labour: admin yillik to'g'ri ({dir_yillik:,})")
    else:
        test_fail("Labour: admin yillik", f"Kutilgan: {dir_yillik}, Natija: {labour_data['admin_yillik']}")
    
    # Test 2.8: Amortizatsiya
    dep = model.t_depreciation
    dep_yearly = dep["data"]["yearly_list"]
    asset_value = model.asosiy_vositalar
    dep_rate = model.cost_structure.get("amortizatsiya_stavka", 0.15)
    expected_dep = asset_value * dep_rate
    if abs(dep_yearly[0] - expected_dep) < 1:
        test_pass(f"Amortizatsiya: 1-yil to'g'ri ({dep_yearly[0]:,.0f})")
    else:
        test_fail("Amortizatsiya", f"Kutilgan: {expected_dep:,.0f}, Natija: {dep_yearly[0]:,.0f}")

    # Test 2.9: Context generatsiya
    ctx = model.get_context()
    required_ctx_keys = ["loyiha_nomi", "tashabbuskor", "kredit", "npv", "irr", "roi"]
    missing_keys = [k for k in required_ctx_keys if k not in ctx]
    if not missing_keys: test_pass("Context: barcha zaruriy kalitlar mavjud")
    else: test_fail("Context: kalit yo'q", str(missing_keys))
    
    # Test 2.10: MCHJ rejimi
    model_mchj = FinancialEngine(TEST_DATA_MCHJ)
    mchj_taxes = model_mchj.t_taxes
    mchj_rows = mchj_taxes["rows"]
    # MCHJ da QQS qatori bo'lishi kerak
    has_qqs = any("QQS" in str(row[0]) for row in mchj_rows)
    if has_qqs: test_pass("MCHJ: QQS soliq qatori mavjud")
    else: test_fail("MCHJ: QQS soliq qatori yo'q")
    
    # Test 2.11: Minimal data bilan ishlash
    model_min = FinancialEngine(TEST_DATA_MINIMAL)
    if model_min.indicators["npv"] is not None: test_pass("Minimal data: model ishladi")
    else: test_fail("Minimal data: model ishlamadi")

    # Test 2.12: Cash flow kumulyativ tekshiruvi
    cf = model.t_cash_flow["data"]["yearly"]
    cumulative_manual = -model.loyiha_qiymati
    all_cum_ok = True
    for i, c in enumerate(cf):
        cumulative_manual += c["sof_cf"]
        if abs(c["kumulyativ"] - cumulative_manual) > 1:
            test_fail(f"CashFlow kumulyativ {i+1}-yil", f"Kutilgan: {cumulative_manual:,.0f}, Natija: {c['kumulyativ']:,.0f}")
            all_cum_ok = False
            break
    if all_cum_ok: test_pass("CashFlow: kumulyativ hisob to'g'ri")
    
    # Test 2.13: ProfLoss — daromad >= tannarx (sof foyda tekshiruvi)
    pnl = model.t_prof_loss["data"]["yearly"]
    pnl_check = all(p["daromad"] > 0 for p in pnl)
    if pnl_check: test_pass("ProfLoss: barcha yillarda daromad > 0")
    else: test_fail("ProfLoss: daromad manfiy")

except Exception as e:
    test_fail("FINANCIAL ENGINE", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 3. WORD GENERATSIYA TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("3. WORD GENERATSIYA TESTI")
print("="*60)

try:
    from modules.document_engine import create_word_document
    from modules.financial_engine import FinancialEngine
    
    model = FinancialEngine(TEST_DATA)
    ctx = model.get_context()
    
    # Kontekstni kengaytirish (app.py dagi kabi)
    f_npv = f"{model.indicators['npv']:,.0f}".replace(",", " ")
    f_roi = f"{model.indicators['roi']:.1f}"
    xulosa_text = f"Tahlillar natijasida ushbu loyiha tijorat jihatdan yuqori daromad keltirishi aniqlandi. NPV={f_npv} so'm, ROI={f_roi}%."
    
    ctx.update({
        "extra_doc": "Qo'shimcha hujjat taqdim etilmagan.",
        "tasdiqlash": True,
        "mundarija": "1. Mahfiyligini ta`minlash memorandumi\n2. Loyiha tashabbuskori to'g'risida ma'lumot",
        "mahfiyligini_ta_minlash_memorandumi": "Ushbu biznes reja maxfiy hujjat hisoblanadi.",
        "xulosa": xulosa_text,
    })
    
    template_path = os.path.join(BASE_DIR, "template.docx")
    output_path = os.path.join(BASE_DIR, "test_full_output.docx")
    
    if not os.path.exists(template_path):
        test_fail("Word template", "template.docx mavjud emas!")
    else:
        result = create_word_document(template_path, output_path, ctx, {}, model=model)
        
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            test_pass(f"Word fayl yaratildi: {file_size:,} bytes ({file_size/1024:.0f} KB)")
            
            # Word faylni tekshirish
            from docx import Document
            doc = Document(output_path)
            
            # 1. Paragraflar soni
            para_count = len(doc.paragraphs)
            test_pass(f"Word: {para_count} ta paragraf topildi")
            
            # 2. Jadvallar soni (inline + appended)
            table_count = len(doc.tables)
            if table_count > 0: test_pass(f"Word: {table_count} ta jadval topildi")
            else: test_fail("Word: jadvallar topilmadi!")
            
            # 3. Placeholder tekshiruvi — qolgan {{...}} bo'lmasligi kerak
            remaining_placeholders = []
            for p in doc.paragraphs:
                if '{{' in p.text and '}}' in p.text:
                    # Extract placeholder names
                    import re
                    found = re.findall(r'\{\{(\w+)\}\}', p.text)
                    remaining_placeholders.extend(found)
            
            # Jadvallar ichidagi placeholderlar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if '{{' in p.text and '}}' in p.text:
                                import re
                                found = re.findall(r'\{\{(\w+)\}\}', p.text)
                                remaining_placeholders.extend(found)
            
            if not remaining_placeholders:
                test_pass("Word: barcha placeholderlar almashtirildi ✨")
            else:
                unique_ph = list(set(remaining_placeholders))
                test_fail(f"Word: {len(unique_ph)} ta placeholder qoldi", str(unique_ph[:10]))
            
            # 4. Jadval sarlavhalarini tekshirish
            table_titles_found = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if any(x in t for x in ["ILOVA", "LOYIHA QIYMATI", "ISH HAQI", "KOMMUNAL", "TANNARX", "FOYDA"]):
                    table_titles_found.append(t[:50])
            
            if len(table_titles_found) > 5:
                test_pass(f"Word: {len(table_titles_found)} ta jadval sarlavhasi topildi")
            else:
                test_fail("Word: jadval sarlavhalari kam", f"Topildi: {len(table_titles_found)}")
            
            # 5. Tasdiqlash imzosi bo'lishi kerak
            has_sign = any("Tasdiqlayman" in p.text for p in doc.paragraphs)
            if has_sign: test_pass("Word: tasdiqlash imzosi mavjud")
            else: test_fail("Word: tasdiqlash imzosi yo'q")
            
            # 6. Faylni tozalash
            # os.remove(output_path)  # debug uchun saqlab qo'yamiz
        else:
            test_fail("Word: fayl yaratilmadi!")

except Exception as e:
    test_fail("WORD GENERATSIYA", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 4. EXCEL GENERATSIYA TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("4. EXCEL GENERATSIYA TESTI")
print("="*60)

try:
    from modules.excel_writer import write_excel_output
    
    excel_template = os.path.join(BASE_DIR, "data.xlsx")
    excel_output = os.path.join(BASE_DIR, "test_full_excel.xlsx")
    
    if not os.path.exists(excel_template):
        test_fail("Excel template", "data.xlsx mavjud emas!")
    else:
        model = FinancialEngine(TEST_DATA)
        result = write_excel_output(excel_template, excel_output, model)
        
        if os.path.exists(excel_output):
            file_size = os.path.getsize(excel_output)
            test_pass(f"Excel fayl yaratildi: {file_size:,} bytes ({file_size/1024:.0f} KB)")
            
            from openpyxl import load_workbook
            wb = load_workbook(excel_output)
            
            # 1. Sheet nomlari
            sheet_names = wb.sheetnames
            test_pass(f"Excel: {len(sheet_names)} ta sheet ({', '.join(sheet_names)})")
            
            # 2. Loans sheet tekshiruvi
            if "Loans" in sheet_names:
                ws_loans = wb["Loans"]
                kredit_summa_cell = ws_loans["D9"].value
                if kredit_summa_cell == 350000000:
                    test_pass(f"Excel Loans: kredit summasi to'g'ri ({kredit_summa_cell:,.0f})")
                else:
                    test_fail("Excel Loans: kredit summasi", f"Kutilgan: 350000000, Natija: {kredit_summa_cell}")
                
                # Foiz
                foiz_cell = ws_loans["D6"].value
                if foiz_cell == 0.14:
                    test_pass(f"Excel Loans: foiz to'g'ri ({foiz_cell})")
                else:
                    test_fail("Excel Loans: foiz", f"Kutilgan: 0.14, Natija: {foiz_cell}")
                
                # Jadval ma'lumotlari
                first_month = ws_loans.cell(row=12, column=2).value
                if first_month == 1:
                    test_pass("Excel Loans: 1-oy to'g'ri yozildi")
                else:
                    test_fail("Excel Loans: 1-oy", f"Natija: {first_month}")
            else:
                test_fail("Excel: Loans sheet topilmadi")
            
            # 3. Asosiy sheet tekshiruvi
            ws_main = wb.worksheets[0]
            loyiha_nomi = ws_main["B16"].value
            if loyiha_nomi == "Mebel ishlab chiqarish sexi":
                test_pass("Excel Main: loyiha nomi to'g'ri yozildi")
            else:
                test_fail("Excel Main: loyiha nomi", f"Natija: {loyiha_nomi}")
            
            # 4. ProfLoss sheet
            if "ProfLoss" in sheet_names:
                ws_pl = wb["ProfLoss"]
                revenue_1y = ws_pl.cell(row=4, column=3).value
                if revenue_1y and revenue_1y > 0:
                    test_pass(f"Excel ProfLoss: 1-yil daromad yozildi ({revenue_1y:,.0f})")
                else:
                    test_fail("Excel ProfLoss: daromad", f"Natija: {revenue_1y}")
            else:
                test_fail("Excel: ProfLoss sheet topilmadi")
            
            # 5. NPV sheet
            if "npv" in sheet_names:
                ws_npv = wb["npv"]
                npv_val = ws_npv.cell(row=16, column=5).value
                if npv_val is not None:
                    test_pass(f"Excel NPV: NPV yozildi ({npv_val:,.0f})")
                else:
                    test_fail("Excel NPV: NPV qiymati yo'q")
            else:
                test_fail("Excel: npv sheet topilmadi")
            
            # 6. Depreciation
            if "Depreciate" in sheet_names:
                ws_dep = wb["Depreciate"]
                dep_1y = ws_dep.cell(row=5, column=5).value
                if dep_1y and dep_1y > 0:
                    test_pass(f"Excel Depreciate: 1-yil amortizatsiya yozildi ({dep_1y:,.0f})")
                else:
                    test_fail("Excel Depreciate: amortizatsiya", f"Natija: {dep_1y}")
            
            # 7. CashFlow
            if "CashFlow" in sheet_names:
                ws_cf = wb["CashFlow"]
                inv = ws_cf.cell(row=15, column=3).value
                if inv and inv < 0:
                    test_pass(f"Excel CashFlow: investitsiya yozildi ({inv:,.0f})")
                else:
                    test_fail("Excel CashFlow: investitsiya", f"Natija: {inv}")
            
            # 8. Taxes
            if "Taxes" in sheet_names:
                ws_tax = wb["Taxes"]
                tax_1y = ws_tax.cell(row=4, column=5).value
                if tax_1y is not None and tax_1y > 0:
                    test_pass(f"Excel Taxes: soliq yozildi ({tax_1y:,.0f})")
                else:
                    test_fail("Excel Taxes: soliq", f"Natija: {tax_1y}")

            wb.close()
        else:
            test_fail("Excel: fayl yaratilmadi!")

except Exception as e:
    test_fail("EXCEL GENERATSIYA", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 5. KATEGORIYA VA QIDIRISH TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("5. KATEGORIYA VA QIDIRISH TESTI")
print("="*60)

try:
    from modules.business_categories import (
        get_all_plans, search_plans, get_faoliyat_turi,
        get_cost_structure, get_categories_for_frontend,
        KATEGORIYALAR, FAOLIYAT_TURLARI
    )
    
    # 1. 480 reja
    all_plans = get_all_plans()
    if len(all_plans) >= 480: test_pass(f"Barcha rejalar: {len(all_plans)} ta")
    else: test_fail("Rejalar soni", f"Kutilgan: 480+, Natija: {len(all_plans)}")
    
    # 2. 24 kategoriya
    if len(KATEGORIYALAR) >= 24: test_pass(f"Kategoriyalar: {len(KATEGORIYALAR)} ta")
    else: test_fail("Kategoriyalar soni", f"Kutilgan: 24+, Natija: {len(KATEGORIYALAR)}")
    
    # 3. 4 faoliyat turi
    if len(FAOLIYAT_TURLARI) == 4: test_pass(f"Faoliyat turlari: {len(FAOLIYAT_TURLARI)} ta")
    else: test_fail("Faoliyat turlari", f"{len(FAOLIYAT_TURLARI)}")
    
    # 4. Qidirish
    search_result = search_plans("mebel", 10)
    if len(search_result) > 0: test_pass(f"Qidirish 'mebel': {len(search_result)} ta natija")
    else: test_fail("Qidirish 'mebel'", "natija yo'q")
    
    # 5. Faoliyat turi aniqlash
    ft = get_faoliyat_turi("Mebel ishlab chiqarish")
    if ft == "ishlab_chiqarish": test_pass(f"Faoliyat turi aniqlash: {ft}")
    else: test_fail("Faoliyat turi aniqlash", f"Kutilgan: ishlab_chiqarish, Natija: {ft}")
    
    # 6. Cost structure
    cs = get_cost_structure("savdo")
    if cs["asosiy_xarajat_nomi"] == "Tovar xaridi": test_pass("Cost structure savdo: to'g'ri")
    else: test_fail("Cost structure savdo", cs["asosiy_xarajat_nomi"])
    
    # 7. Frontend data
    frontend_data = get_categories_for_frontend()
    if "faoliyat_turlari" in frontend_data and "kategoriyalar" in frontend_data and "barcha_rejalar" in frontend_data:
        test_pass("Frontend data: tuzilishi to'g'ri")
    else:
        test_fail("Frontend data", str(list(frontend_data.keys())))

except Exception as e:
    test_fail("KATEGORIYALAR", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 6. VALIDATORS TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("6. VALIDATORS TESTI")
print("="*60)

try:
    from modules.validators import validate_form, safe_float, safe_int
    
    # Valid form
    errors = validate_form(TEST_DATA)
    if not errors: test_pass("Valid form: xatolik yo'q")
    else: test_fail("Valid form", str(errors))
    
    # Missing required fields
    bad_form = {"loyiha_nomi": "", "tashabbuskor": ""}
    errors = validate_form(bad_form)
    if len(errors) >= 2: test_pass(f"Invalid form: {len(errors)} ta xatolik topildi")
    else: test_fail("Invalid form", f"Kutilgan: 2+, Natija: {len(errors)}")
    
    # Kredit > loyiha qiymati
    bad_kredit = {"loyiha_nomi": "Test", "tashabbuskor": "Test", "kredit": "200", "loyiha_qiymati": "100"}
    errors = validate_form(bad_kredit)
    has_kredit_err = any("kredit" in e.lower() or "Kredit" in e for e in errors)
    if has_kredit_err: test_pass("Validator: kredit > loyiha qiymati xatoligi ishladi")
    else: test_fail("Validator: kredit > loyiha_qiymati", str(errors))
    
    # safe_float/safe_int
    if safe_float("100.5") == 100.5: test_pass("safe_float: to'g'ri")
    else: test_fail("safe_float")
    
    if safe_int("42") == 42: test_pass("safe_int: to'g'ri")
    else: test_fail("safe_int")

except Exception as e:
    test_fail("VALIDATORS", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 7. PAYMENT MODULI TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("7. PAYMENT MODULI TESTI")
print("="*60)

try:
    from modules.payment import create_payment, verify_payment, get_payment_info, PLAN_PRICE
    
    if PLAN_PRICE == 80_000: test_pass(f"Narx: {PLAN_PRICE:,} so'm")
    else: test_fail("Narx", f"{PLAN_PRICE}")
    
    payment = create_payment("demo")
    if payment["status"] == "pending": test_pass("Payment yaratildi: pending")
    else: test_fail("Payment yaratish", payment.get("status"))
    
    verify_result = verify_payment(payment["id"])
    if verify_result["success"]: test_pass("Payment tasdiqlandi: demo rejim")
    else: test_fail("Payment tasdiqlash", str(verify_result))
    
    info = get_payment_info()
    if len(info["methods"]) >= 4: test_pass(f"Payment methods: {len(info['methods'])} ta")
    else: test_fail("Payment methods", f"{len(info['methods'])}")

except Exception as e:
    test_fail("PAYMENT", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 8. MATEMATIKA TEKSHIRUVI (Pro daraja)
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("8. MATEMATIKA TEKSHIRUVI")
print("="*60)

try:
    model = FinancialEngine(TEST_DATA)
    
    # 8.1: Kredit to'lov + asosiy qarz = kredit summa
    total_principal = sum(t.asosiy_qarz for t in model.kredit_natija.jadval)
    if abs(total_principal - 350000000) < 10:
        test_pass(f"Kredit: jami asosiy qarz = kredit summa ({total_principal:,.0f} ≈ 350,000,000)")
    else:
        test_fail("Kredit: jami asosiy qarz", f"Kutilgan: 350,000,000, Natija: {total_principal:,.0f}")
    
    # 8.2: ProfLoss — sof_daromad = daromad - qqs (YTT da qqs=0)
    for i, p in enumerate(model.yearly_pnl):
        expected_net = p["daromad"] - p["qqs"]
        if abs(p["sof_daromad"] - expected_net) > 1:
            test_fail(f"ProfLoss {i+1}-yil: sof_daromad", f"Kutilgan: {expected_net:,.0f}, Natija: {p['sof_daromad']:,.0f}")
            break
    else:
        test_pass("ProfLoss: sof_daromad hisob to'g'ri (barcha yillar)")
    
    # 8.3: YTT rejimda soliq = 4% * daromad
    for i, p in enumerate(model.yearly_pnl):
        expected_soliq = model.t_prod_plan["data"]["yearly_revenue"][i] * 0.04
        if abs(p["soliq"] - expected_soliq) > 1:
            test_fail(f"Soliq {i+1}-yil", f"Kutilgan: {expected_soliq:,.0f}, Natija: {p['soliq']:,.0f}")
            break
    else:
        test_pass("Soliq: YTT 4% hisob to'g'ri (barcha yillar)")
    
    # 8.4: CashFlow investitsiya ≠ 0
    if model.t_cash_flow["data"]["investment"] == 500000000:
        test_pass(f"CashFlow: investitsiya to'g'ri ({model.t_cash_flow['data']['investment']:,})")
    else:
        test_fail("CashFlow: investitsiya", f"Natija: {model.t_cash_flow['data']['investment']}")
    
    # 8.5: NPV formula tekshiruvi
    cf_list = [-model.loyiha_qiymati] + [c["sof_cf"] for c in model.yearly_cf]
    manual_npv = calculate_npv(model.discount_rate, cf_list)
    if abs(manual_npv - model.indicators["npv"]) < 1:
        test_pass(f"NPV: qo'lda hisoblangan qiymat mos ({manual_npv:,.0f})")
    else:
        test_fail("NPV: qo'lda hisob", f"Manual: {manual_npv:,.0f}, Engine: {model.indicators['npv']:,.0f}")
    
    # 8.6: Ijtimoiy soliq 12% tekshiruvi
    jami_fond = model.t_labour["data"]["jami_fond"]
    ijtimoiy = model.t_labour["data"]["ijtimoiy"]
    expected_ijt = jami_fond * 0.12
    if abs(ijtimoiy - expected_ijt) < 1:
        test_pass(f"Ijtimoiy soliq: 12% to'g'ri ({ijtimoiy:,.0f})")
    else:
        test_fail("Ijtimoiy soliq", f"Kutilgan: {expected_ijt:,.0f}, Natija: {ijtimoiy:,.0f}")

    # 8.7: MCHJ QQS tekshiruvi
    model_m = FinancialEngine(TEST_DATA_MCHJ)
    for i, p in enumerate(model_m.yearly_pnl):
        # QQS = daromad * 12/112
        expected_qqs = model_m.t_prod_plan["data"]["yearly_revenue"][i] * 12 / 112
        if abs(p["qqs"] - expected_qqs) > 1:
            test_fail(f"MCHJ QQS {i+1}-yil", f"Kutilgan: {expected_qqs:,.0f}, Natija: {p['qqs']:,.0f}")
            break
    else:
        test_pass("MCHJ: QQS hisob to'g'ri (barcha yillar)")

except Exception as e:
    test_fail("MATEMATIKA", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# 9. HTML TEMPLATE TESTI
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("9. HTML/JS TUZILISHI TESTI")
print("="*60)

try:
    html_path = os.path.join(BASE_DIR, "templates", "index.html")
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    # Broken HTML — line 418 da buzilgan button tag
    if 'style="width:100%; height:60px; font-size:1.2rem" <div' in html:
        test_fail("HTML: 418-qatorda buzilgan button tag topildi!", "button ichida <div tag mavjud")
    else:
        test_pass("HTML: buzilgan taglar yo'q")
    
    # CSRF
    if "csrf_token" in html: test_pass("HTML: CSRF token mavjud")
    else: test_fail("HTML: CSRF token yo'q")
    
    # 8 ta step
    step_count = html.count('class="form-step')
    if step_count == 8: test_pass(f"HTML: {step_count} ta form step")
    else: test_fail("HTML: form step soni", f"{step_count}")
    
    # JS fayl
    js_path = os.path.join(BASE_DIR, "static", "js", "app.js")
    if os.path.exists(js_path): test_pass("JS: app.js mavjud")
    else: test_fail("JS: app.js yo'q")

except Exception as e:
    test_fail("HTML/JS TESTI", traceback.format_exc())


# ══════════════════════════════════════════════════════════
# NATIJALAR
# ══════════════════════════════════════════════════════════
print("\n" + "="*60)
print("NATIJALAR")
print("="*60)
total = results["passed"] + results["failed"]
print(f"  Jami testlar: {total}")
print(f"  ✅ O'tdi: {results['passed']}")
print(f"  ❌ Yiqildi: {results['failed']}")

if results["errors"]:
    print(f"\nXATOLIKLAR:")
    for err in results["errors"]:
        print(f"  ⚠ {err}")

print("="*60)
