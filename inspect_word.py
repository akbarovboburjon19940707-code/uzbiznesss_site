import docx

doc = docx.Document("e:/uzbiznesss_site/template.docx")
print("Tables in Word:")
for i, table in enumerate(doc.tables):
    print(f"Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
    # Read the first row to get an idea
    try:
        row_text = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
         # if it is long, truncate
        row_text = [r[:30] for r in row_text]
        print(f"  Header: {row_text}")
    except:
        pass
        
print("Searching for known Excel table headers/placeholders...")
for p in doc.paragraphs:
    if "{{" in p.text:
         print("Found var in para:", p.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
             for p in cell.paragraphs:
                  if "{{" in p.text:
                       # just print a sample to avoid huge output
                       pass

