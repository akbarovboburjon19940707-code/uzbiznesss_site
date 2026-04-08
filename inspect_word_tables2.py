import docx

doc = docx.Document("e:/uzbiznesss_site/template.docx")
tables_to_check = [7, 8, 9, 10] # 0-indexed, so 8th to 11th table
for i in tables_to_check:
    if i < len(doc.tables):
        table = doc.tables[i]
        print(f"\n--- Table {i+1} ---")
        for r_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.replace('\n', ' ').strip()[:50])
            print(f"Row {r_idx+1}:", " | ".join(row_text))

