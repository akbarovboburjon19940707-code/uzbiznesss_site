
import docx
import os

template_path = "d:/uzbiznesss_site/template.docx"
if not os.path.exists(template_path):
    print(f"Template not found at {template_path}")
    exit()

doc = docx.Document(template_path)

search_text = "6.3. Kommunikatsiya va infratuzilma xarajatlari"
found_para = False

print(f"Searching for '{search_text}' in paragraphs...")
for i, p in enumerate(doc.paragraphs):
    if search_text in p.text:
        print(f"Found at paragraph index {i}: {p.text}")
        found_para = True
        # Print following paragraphs too
        for next_p in doc.paragraphs[i+1:i+5]:
            print(f"  Next: {next_p.text}")
        break

if not found_para:
    print("Not found in paragraphs, searching in tables...")
    for i, t in enumerate(doc.tables):
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                if search_text in cell.text:
                    print(f"Found in Table {i}, Row {r_idx}, Cell {c_idx}: {cell.text}")
                    found_para = True

if not found_para:
    print("Searching for similar text (6.3)...")
    for i, p in enumerate(doc.paragraphs):
        if "6.3" in p.text:
            print(f"Found '6.3': {p.text}")
