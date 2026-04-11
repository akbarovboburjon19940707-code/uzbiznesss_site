
import os
import logging
from modules.document_engine import convert_to_pdf
from docx import Document

logging.basicConfig(level=logging.INFO)

def test_conversion():
    input_docx = "test_debug.docx"
    output_pdf = "test_debug.pdf"
    
    # Create a simple docx
    doc = Document()
    doc.add_heading('Test PDF Generation', 0)
    doc.add_paragraph('This is a test document to debug PDF conversion.')
    doc.save(input_docx)
    
    print(f"File created: {input_docx}")
    
    try:
        result = convert_to_pdf(input_docx, output_pdf)
        if result:
            print(f"SUCCESS: PDF generated at {result}")
        else:
            print("FAILURE: PDF generation returned None")
    except Exception as e:
        print(f"ERROR: Exception during conversion: {e}")
        import traceback
        traceback.print_exc()

Saytda PDF hosil bo'lmasligining bir necha asosiy sabablari bo'lishi mumkin. Hozirgi kod tahlili va o'tkazilgan testlar natijasida quyidagi muammolarni aniqladim:

Muhit farqi (Linux vs Windows): Agar saytni Render, Heroku yoki boshqa Linux serverga joylagan bo'lsangiz, docx2pdf (Microsoft Word orqali ishlashi sababli) u yerda ishlamaydi. Serverda LibreOffice o'rnatilgan bo'lishi shart.
Kutubxonalar yetishmasligi: requirements.txt faylida PDF yaratish uchun kerakli docx2pdf va pywin32 kutubxonalari ko'rsatilmagan.
Word bloklanishi: Windows'da bir vaqtning o'zida bir nechta foydalanuvchi PDF so'rasa, Microsoft Word qotib qolishi yoki xatolik berishi mumkin (uni faqat bitta potokda ishlatish kerak).
Ushbu muammolarni bartaraf etish uchun 
implementation_plan.md
 rejasini tayyorladim. Rejani ko'rib chiqib tasdiqlasangiz, tuzatish ishlarini boshlayman.

Sizga savollarim:

Sayt hozir qayerda ishlayapti (o'zingizni kompyuteringizdami yoki serverdami)?
Agar serverda bo'lsa, u Linuxmi yoki Windows?
Kompyuteringizda Microsoft Word dasturi o'rnatilganmi?
Implementation Plan
1
test_pdf.py
Failed to move changes
failed to checkout worktree changes: main branch has uncommitted changes; please commit or stash your changes before checking out the worktree


if __name__ == "__main__":
    test_conversion()
