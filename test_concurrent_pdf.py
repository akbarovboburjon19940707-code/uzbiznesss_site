
import threading
import time
import os
from modules.document_engine import convert_to_pdf, create_word_document
from docx import Document

def worker(worker_id):
    input_docx = f"test_concurrent_{worker_id}.docx"
    output_pdf = f"test_concurrent_{worker_id}.pdf"
    
    # Create simple docx
    doc = Document()
    doc.add_heading(f'Concurrent Test {worker_id}', 0)
    doc.add_paragraph(f'Worker {worker_id} is converting this document.')
    doc.save(input_docx)
    
    print(f"Worker {worker_id}: Starting conversion...")
    start = time.time()
    result = convert_to_pdf(input_docx, output_pdf)
    end = time.time()
    
    if result:
        print(f"Worker {worker_id}: SUCCESS in {end-start:.2f}s")
    else:
        print(f"Worker {worker_id}: FAILURE")

if __name__ == "__main__":
    threads = []
    # Test with 3 concurrent workers
    for i in range(3):
        t = threading.Thread(target=worker, args=(i,))
        threads.append(t)
        t.start()
        
    for t in threads:
        t.join()
    
    print("All workers finished.")
